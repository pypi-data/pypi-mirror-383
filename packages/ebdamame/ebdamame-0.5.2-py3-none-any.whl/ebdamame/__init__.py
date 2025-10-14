"""
Contains high level functions to process .docx files
"""

import gc
import itertools
import logging
import re
import sys
from io import BytesIO
from pathlib import Path
from typing import Generator, Iterable, Optional, Union

import attrs
import docx
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

_logger = logging.getLogger(__name__)

_is_python_version_314 = sys.version_info[0:2] == (3, 14)
_is_manually_triggered_garbage_collection_required = _is_python_version_314

"""
I don't know the reason why, but the CI failed in Python 3.14 with the following message:
"Error: Process completed with exit code 143."
based on the commit https://github.com/Hochfrequenz/ebdamame/pull/363/commits/b6a456345d46a11fe09c6c1c32ff66e62cb1392c

The python-docx repo as of 2025-10-13 mentions one open issue which might be related:
https://github.com/python-openxml/python-docx/issues/1428
Also in the CPython repository there is an open regression bug, that maybe affects ebdamame internally:
https://github.com/python/cpython/issues/139951
So as a workaround, we trigger garbage collection manually after working with a docx file.
"""


def get_document(docx_file_path: Path) -> DocumentType:
    """
    opens and returns the document specified in the docx_file_path using python-docx
    """
    with open(docx_file_path, "rb") as docx_file:
        source_stream = BytesIO(docx_file.read())
        # Originally I tried the recipe from
        # https://python-docx.readthedocs.io/en/latest/user/documents.html#opening-a-file-like-document
        # but then switched from StringIO to BytesIO (without explicit 'utf-8') because of:
        # UnicodeDecodeError: 'charmap' codec can't decode byte 0x81 in position 605: character maps to <undefined>
    try:
        document = docx.Document(source_stream)
        _logger.info("Successfully read the file '%s'", docx_file_path)
        return document
    finally:
        source_stream.close()


def _get_tables_and_paragraphs(document: DocumentType) -> Generator[Union[Table, Paragraph], None, None]:
    """
    Yields tables and paragraphs from the given document in the order in which they occur in the document.
    This is helpful because document.tables and document.paragraphs are de-coupled and give you no information which
    paragraph follows which table.
    """
    parent_elements = document.element.body
    for item in parent_elements.iterchildren():
        if isinstance(item, CT_P):
            yield Paragraph(item, document)
        elif isinstance(item, CT_Tbl):
            yield Table(item, document)
        else:
            _logger.debug("Item %s is neither Paragraph nor Table", str(item))


_ebd_key_pattern = re.compile(r"^E_\d{4}$")
_ebd_key_with_heading_pattern = re.compile(r"^(?P<key>E_\d{4})_?(?P<title>.*)\s*$")


class TableNotFoundError(Exception):
    """
    an error that is raised when a requested table was not found
    """

    def __init__(self, ebd_key: str):
        self.ebd_key = ebd_key


_ebd_cell_pattern = re.compile(r"^(?:ja|nein)\s*(?:Ende|\d+)$")
"""
any EBD table shall contain at least one cell that matches this pattern
"""


def _cell_is_probably_from_an_ebd_cell(cell: _Cell) -> bool:
    if "" in cell.text:
        return True
    if cell.text in {"ja", "nein"}:
        return True
    if "à" in cell.text:
        # the rightarrow in wrong encoding
        return True
    if _ebd_cell_pattern.match(cell.text):
        return True
    if cell.text.strip().startswith("Cluster:") or cell.text.startswith("Hinweis:"):
        return True
    return False


def _table_is_an_ebd_table(table: Table) -> bool:
    """
    Returns true iff the table "looks like" an EB-Table.
    This is to distinguish between tables that are inside the same subsection that describes an EBD but are not part
    of the decision tree at all (e.g. in E_0406 the tables about Artikel-IDs).
    """
    for row in table.rows:
        try:
            for cell in row.cells:
                if _cell_is_probably_from_an_ebd_cell(cell):
                    return True
        except IndexError:  # don't ask me why this happens; It's the internals of python-docx
            continue
    return False


def _table_is_first_ebd_table(table: Table) -> bool:
    """
    Returns true if the first row of a table contains "Prüfende Rolle".
    We assume that each EBD table has a header row with
    "Prüfende Rolle" in the first column.
    """
    return "prüfende rolle" in table.rows[0].cells[0].text.lower()


@attrs.define(kw_only=True, frozen=True)
class EbdNoTableSection:
    """
    Represents an empty section in the document
    """

    ebd_key: str = attrs.field(validator=attrs.validators.instance_of(str))
    remark: str = attrs.field(validator=attrs.validators.instance_of(str))


# pylint:disable=too-many-branches
def is_heading(paragraph: Paragraph) -> bool:
    """
    Returns True if the paragraph is a heading.
    """
    return paragraph.style is not None and paragraph.style.style_id in {
        "berschrift1",
        "berschrift2",
        "berschrift3",
    }


def get_ebd_docx_tables(docx_file_path: Path, ebd_key: str) -> list[Table] | EbdNoTableSection:
    """
    Opens the file specified in `docx_file_path` and returns the tables that relate to the given `ebd_key`.

    This function processes the document to find tables associated with the given `ebd_key`.
    There might be more than one table for a single EBD table due to inconsistencies and manual editing during
    the creation of the documents by EDI@Energy.
    There are sections relating to the EBD key without any tables.
    In this case, the section is identified and the related paragraph is captured as a remark
    (e.g. 'Es ist das EBD E_0556 zu nutzen.' for EBD_0561).

    Args:
        docx_file_path (Path): The path to the .docx file to be processed.
        ebd_key (str): The EBD key to search for in the document.

    Returns:
        list[Table] | EbdNoTableSection: A list of `Table` objects if tables are found, or an `EbdNoTableSection` object
        if no tables are found but the section is identified and are remark is captured.

    Raises:
        TableNotFoundError: If no tables related to the given `ebd_key` are found in the document.
    """
    if _ebd_key_pattern.match(ebd_key) is None:
        raise ValueError(f"The ebd_key '{ebd_key}' does not match {_ebd_key_pattern.pattern}")
    document = get_document(docx_file_path)

    empty_ebd_text: str | None = None  # paragraph text if there is no ebd table
    found_table_in_subsection: bool = False
    is_inside_subsection_of_requested_table: bool = False
    tables: list[Table] = []
    tables_and_paragraphs = _get_tables_and_paragraphs(document)
    for table_or_paragraph in tables_and_paragraphs:
        if isinstance(table_or_paragraph, Paragraph):
            paragraph: Paragraph = table_or_paragraph
            # Assumptions:
            # 1. before each EbdTable there is a paragraph whose text starts with the respective EBD key
            # 2. there are no duplicates
            is_ebd_heading_of_requested_ebd_key = paragraph.text.startswith(ebd_key)
            if is_inside_subsection_of_requested_table and is_heading(paragraph):
                _logger.warning("No EBD table found in subsection for: '%s'", ebd_key)
                break
            if is_inside_subsection_of_requested_table and paragraph.text.strip() != "":
                if empty_ebd_text is None:
                    # the first text paragraph after we found the correct section containing the ebd key
                    empty_ebd_text = paragraph.text.strip()
                else:
                    empty_ebd_text += ("\n") + paragraph.text.strip()
            is_inside_subsection_of_requested_table = (
                is_ebd_heading_of_requested_ebd_key or is_inside_subsection_of_requested_table
            )
        if isinstance(table_or_paragraph, Table) and is_inside_subsection_of_requested_table:
            found_table_in_subsection = True
        if (
            isinstance(table_or_paragraph, Table)
            and is_inside_subsection_of_requested_table
            and _table_is_an_ebd_table(table_or_paragraph)
            and _table_is_first_ebd_table(table_or_paragraph)
        ):
            table: Table = table_or_paragraph
            tables.append(table)
            # Now we have to check if the EBD table spans multiple pages, and _maybe_ we have to collect more tables.
            # The funny thing is: Sometimes the authors create multiple tables split over multiple lines which belong
            # together, sometimes they create 1 proper table that spans multiple pages.
            # The latter case (1 docx table spanning >1 pages) is transparent to the extraction logic; i.e. python-docx
            # treats a single table that spans multiple pages just the same as a table on only 1 page.
            for next_item in tables_and_paragraphs:  # start iterating from where the outer loop paused
                if isinstance(next_item, Table):
                    # this is the case that the authors created multiple single tables on single adjacent pages
                    # if table_is_an_ebd_table(table):
                    if _table_is_an_ebd_table(next_item):
                        tables.append(next_item)
                elif isinstance(next_item, Paragraph):
                    if next_item.text.startswith("S_") or next_item.text.startswith("E_"):
                        # this is the case that the authors created 1 table that spans multiple pages
                        # and we're done collecting tables for this EBD key
                        break
                    continue
                else:
                    break  # inner loop because if no other table will follow
                    # we're done collecting the tables for this EBD key
        if is_inside_subsection_of_requested_table and len(tables) > 0:  # this means: we found the table
            # break the outer loop, too; no need to iterate any further
            break
    if not any(tables):
        if not is_inside_subsection_of_requested_table:
            raise TableNotFoundError(ebd_key=ebd_key)
        if empty_ebd_text is None:
            if found_table_in_subsection:
                # probably there is an error while scraping the tables
                raise TableNotFoundError(ebd_key=ebd_key)
            return EbdNoTableSection(ebd_key=ebd_key, remark="")
        return EbdNoTableSection(ebd_key=ebd_key, remark=empty_ebd_text.strip())
    try:
        return tables
    finally:
        if _is_manually_triggered_garbage_collection_required:
            del document
            gc.collect()


# pylint:disable=too-few-public-methods
@attrs.define(kw_only=True, frozen=True)
class EbdChapterInformation:
    """
    Contains information about where an EBD is located within the document.
    If the heading is e.g. "5.2.1" we denote this as:
    * chapter 5
    * section 2
    * subsection 1
    """

    chapter: int = attrs.field(
        validator=attrs.validators.and_(attrs.validators.instance_of(int), attrs.validators.ge(1))
    )
    chapter_title: Optional[str] = attrs.field(validator=attrs.validators.optional(attrs.validators.instance_of(str)))
    section: int = attrs.field(
        validator=attrs.validators.and_(attrs.validators.instance_of(int), attrs.validators.ge(1))
    )

    section_title: Optional[str] = attrs.field(validator=attrs.validators.optional(attrs.validators.instance_of(str)))
    subsection: int = attrs.field(
        validator=attrs.validators.and_(attrs.validators.instance_of(int), attrs.validators.ge(1))
    )

    subsection_title: Optional[str] = attrs.field(
        validator=attrs.validators.optional(attrs.validators.instance_of(str))
    )


def _enrich_paragraphs_with_sections(
    paragraphs: Iterable[Paragraph],
) -> Generator[tuple[Paragraph, EbdChapterInformation], None, None]:
    """
    Yield each paragraph + the "Kapitel" in which it is found.
    """
    chapter_counter = itertools.count(start=1)
    chapter = 1
    chapter_title: Optional[str] = None
    section_counter = itertools.count(start=1)
    section = 1
    section_title: Optional[str] = None
    subsection_counter = itertools.count(start=1)
    subsection = 1
    subsection_title: Optional[str] = None
    for paragraph in paragraphs:
        # since pyton-docx 1.1.2 there are type hints; seems like the style is not guaranteed to be not None
        match paragraph.style.style_id:  #  type:ignore[union-attr]
            case "berschrift1":
                chapter = next(chapter_counter)
                chapter_title = paragraph.text.strip()
                section_counter = itertools.count(start=1)
                section_title = None
                subsection_counter = itertools.count(start=1)
                subsection_title = None
            case "berschrift2":
                section = next(section_counter)
                section_title = paragraph.text.strip()
                subsection_counter = itertools.count(start=1)
                subsection_title = None
            case "berschrift3":
                subsection = next(subsection_counter)
                subsection_title = paragraph.text.strip()
        location = EbdChapterInformation(
            chapter=chapter,
            section=section,
            subsection=subsection,
            chapter_title=chapter_title,
            section_title=section_title,
            subsection_title=subsection_title,
        )
        _logger.debug("Handling Paragraph %i.%i.%i", chapter, section, subsection)
        yield paragraph, location


def get_all_ebd_keys(docx_file_path: Path) -> dict[str, tuple[str, EbdChapterInformation]]:
    """
    Extract all EBD keys from the given file.
    Returns a dictionary with all EBD keys as keys and the respective EBD titles as values.
    E.g. key: "E_0003", value: "Bestellung der Aggregationsebene RZ prüfen"
    """
    document = get_document(docx_file_path)
    result: dict[str, tuple[str, EbdChapterInformation]] = {}
    for paragraph, ebd_kapitel in _enrich_paragraphs_with_sections(document.paragraphs):
        match = _ebd_key_with_heading_pattern.match(paragraph.text)
        if match is None:
            contains_ebd_number = paragraph.text.lstrip().startswith("E_")
            if contains_ebd_number:
                _logger.warning("Found EBD number but could not match: '%s'", paragraph.text)
            continue
        ebd_key = match.groupdict()["key"]
        title = match.groupdict()["title"]
        result[ebd_key] = (title, ebd_kapitel)
        _logger.debug("Found EBD %s: '%s' (%s)", ebd_key, title, ebd_kapitel)
    _logger.info("%i EBD keys have been found", len(result))
    try:
        return result
    finally:
        if _is_manually_triggered_garbage_collection_required:
            del document
            gc.collect()
