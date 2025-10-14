# !/usr/bin/python3
# -*- coding:utf-8 -*-
"""
@Author: xiaodong.li
@Time: 7/28/2025 1:12 PM
@Description: Description
@File: word_replacer.py
"""
from typing import Dict

from docx import Document

from edetek_words.common.docx_utils import parse_text_segments, apply_text_segments_to_runs
from edetek_words.common.word_cleaner import clean_data
from edetek_words.dto.segment_dto import SegmentDTO


class WordReplacer:
    def __init__(self, doc_path):
        self.doc = Document(doc_path)
        self.doc_path = doc_path

    @staticmethod
    def _replace_text_in_runs(runs, segment: SegmentDTO):
        full_text = ''.join(run.text for run in runs)
        if segment.original_text not in full_text:
            full_text = clean_data(full_text)
            if segment.original_text not in full_text:
                return False
        replaced_text = full_text.replace(segment.original_text, segment.translated_dto.translated_text, 1)
        replaced_text = clean_data(replaced_text)
        first_valid_run_idx = None
        for idx, run in enumerate(runs):
            if clean_data(run.text):
                first_valid_run_idx = idx
                break
        if first_valid_run_idx is None:
            raise Exception("first_valid_run_idx is None")
        for idx, run in enumerate(runs):
            if idx >= first_valid_run_idx:
                run.text = ""
                run.bold = False
        if segment.styled_segments:
            text_segments = parse_text_segments(segment.translated_dto.translated_text,
                                                [s.text for s in segment.translated_dto.styled_segments])
            apply_text_segments_to_runs(runs, first_valid_run_idx, text_segments)
        else:
            first_run = runs[first_valid_run_idx]
            first_run.text = replaced_text
            first_run.bold = False
        return True

    def replace(self, replace_dict: Dict[str, SegmentDTO]):
        for para in self.doc.paragraphs:
            self._replace_in_paragraph(para, replace_dict)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replace_dict)
        for section in self.doc.sections:
            header = section.header
            for para in header.paragraphs:
                self._replace_in_paragraph(para, replace_dict)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_in_paragraph(para, replace_dict)
            footer = section.footer
            for para in footer.paragraphs:
                self._replace_in_paragraph(para, replace_dict)
        return self

    def _replace_in_paragraph(self, paragraph, replace_dict):
        clean_para_text = clean_data(paragraph.text)
        if clean_para_text in replace_dict:
            self._replace_text_in_runs(paragraph.runs, replace_dict[clean_para_text])
            return

        for old_text, segment in replace_dict.items():
            if old_text in clean_para_text:
                self._replace_text_in_runs(paragraph.runs, segment)

    def save(self):
        self.doc.save(self.doc_path)
