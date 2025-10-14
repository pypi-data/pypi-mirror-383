# !/usr/bin/python3
# -*- coding:utf-8 -*-
"""
@Author: xiaodong.li
@Time: 7/28/2025 4:53 PM
@Description: Description
@File: word_extractor.py
"""
import re
from pathlib import Path
from typing import List, Set, Optional

from docx import Document

from edetek_words.common.docx_utils import get_bold_semantic_texts, split_runs_by_tab
from edetek_words.common.json_utils import save_json
from edetek_words.common.path import output_doc_path
from edetek_words.common.word_cleaner import clean_data
from edetek_words.dto.segment_dto import SegmentDTO
from edetek_words.dto.styled_text_segment import StyledTextSegment, FontStyle


def exclude_text(text):
    if re.fullmatch(r"\d+", text):
        return True
    if re.fullmatch(r"\d{1,2}:\d{2}", text):
        return True
    return False


class WordExtractor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self.segments: List[SegmentDTO] = []
        self._seen: Set[str] = set()
        self._deduplicate = True

    def extract(self, save_json_file: bool = False, deduplicate: bool = True) -> List[SegmentDTO]:
        self._deduplicate = deduplicate
        self._process_paragraphs(self.doc.paragraphs)
        self._process_tables(self.doc.tables)
        self._process_sections()
        if save_json_file:
            filename = Path(self.file_path).with_suffix(".json").name
            filepath = output_doc_path(filename)
            if filepath.exists() and filepath.is_file():
                filepath.unlink()
            save_json(str(filepath), [segment.original_text for segment in self.segments],
                      ensure_ascii=False)  # for debug
        return self.segments

    def _process_paragraphs(self, paragraphs):
        for para in paragraphs:
            self._handle_paragraph(para)

    def _process_tables(self, tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self._process_paragraphs(cell.paragraphs)

    def _process_sections(self):
        for section in self.doc.sections:
            self._process_paragraphs(section.header.paragraphs)
            self._process_tables(section.header.tables)
            self._process_paragraphs(section.footer.paragraphs)

    def _handle_paragraph(self, para):
        cleaned_text = clean_data(para.text)
        if not cleaned_text:
            return
        processed_segments = []
        for line in re.split(r'[\n\t]+', cleaned_text):
            cleaned_line = clean_data(line)
            if not cleaned_line or (self._deduplicate and cleaned_line in self._seen):
                return
            if exclude_text(cleaned_line):
                return
            segment = SegmentDTO(
                original_text=cleaned_line,
            )
            processed_segments.append(segment)
            self.segments.append(segment)
            self._seen.add(cleaned_line)
        if not processed_segments:
            return
        self._process_bold_styles(para.runs, processed_segments)

    @staticmethod
    def _find_segment_by_original_text(original_text, segments) -> Optional[SegmentDTO]:
        for segment in segments:
            if segment.original_text == original_text:
                return segment
        return None

    def _process_bold_styles(self, runs, processed_segments):
        """处理粗体样式的独立方法"""
        flag = self._assign_bold_segments_to_groups([runs], processed_segments, raise_on_missing=False)
        if flag:
            return
        group_runs = split_runs_by_tab(runs)
        self._assign_bold_segments_to_groups(group_runs, processed_segments)

    def _assign_bold_segments_to_groups(self, runs_groups, segments: List[SegmentDTO], raise_on_missing: bool = True):
        """
        为每组 Run 分配加粗文本到对应的 SegmentDTO 对象。
        """
        for runs_group in runs_groups:
            bold_texts = get_bold_semantic_texts(runs_group)
            if not bold_texts:
                continue
            cleaned_text = clean_data("".join(run.text for run in runs_group))
            if not cleaned_text:
                continue
            segment = self._find_segment_by_original_text(cleaned_text, segments)
            if segment:
                segment.styled_segments = [StyledTextSegment(text, [FontStyle.BOLD]) for text in bold_texts]
            elif raise_on_missing:
                raise Exception("An error occurred while parsing the file. Please contact the administrator.")
            else:
                return False
        return True
