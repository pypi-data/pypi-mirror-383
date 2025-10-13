import logging
import tempfile
from typing import Callable, List

import pandas as pd
import plotly.graph_objects as go  # type: ignore
from buildingspy.io.outputfile import Reader  # type: ignore
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from matplotlib import pyplot as plt
from matplotlib.figure import Figure as pyFigure
from plotly.graph_objects import Figure as plotlyFigure
from plotly.subplots import make_subplots  # type: ignore

from trano.elements import BaseElement, NamedFigure

logger = logging.getLogger(__name__)
FIGURE_COUNT = 1


def plot(data: Reader, figure: NamedFigure, show: bool = True) -> pyFigure:
    fig, ax = plt.subplots(figsize=(10, 6))
    fig.subplots_adjust(right=0.9)
    twin1 = ax.twinx()
    plots = []
    tkw = {"size": 4, "width": 1.5}
    for axis, figure_axis in [(ax, figure.left_axis), (twin1, figure.right_axis)]:
        for line in figure_axis.lines:
            try:
                variable = data.varNames(f"{line.key}$")[0]
                line_data = pd.DataFrame(data.values(variable))
            except (KeyError, IndexError):
                logger.warning(f"Key {line.key} not found in data")
                continue
            (p,) = axis.plot(
                line_data.loc[1],
                linestyle=line.line_style,
                linewidth=line.line_width,
                color=line.color,
                label=line.label,
            )
            plots.append(p)
            axis.set_ylabel(figure_axis.label)
            axis.yaxis.label.set_color(p.get_color())
            axis.tick_params(axis="y", colors=p.get_color(), **tkw)

    ax.set_xlabel("Simulation time [-]")
    ax.tick_params(axis="x", **tkw)
    ax.legend(handles=plots)
    plt.xticks(rotation=45)  # Rotate the tick labels
    plt.tight_layout()
    if show:
        plt.show()
    return fig


def plot_plot_ly_many(
    data: Reader, figures: List[NamedFigure], show: bool = False
) -> plotlyFigure:

    fig = make_subplots(specs=[[{"secondary_y": True}]])
    for figure in figures:
        for axis, figure_axis in enumerate([figure.left_axis, figure.right_axis]):
            for line in figure_axis.lines:
                try:
                    variable = data.varNames(f"{line.key}$")[0]
                    line_data = pd.DataFrame(data.values(variable))
                except (KeyError, IndexError):
                    logger.warning(f"Key {line.key} not found in data")
                    continue

                if line_data.empty:
                    continue

                fig.add_trace(
                    go.Scatter(
                        x=line_data.loc[0],
                        y=line_data.loc[1],
                        mode="lines",
                        name=f"{figure.name} - {line.label}",
                    ),
                    secondary_y=bool(axis),
                )
    if not fig.data:
        return None
    fig.update_layout(
        xaxis_title="Simulation time [-]",
        yaxis_title=figure.left_axis.label,
        yaxis2_title=figure.right_axis.label,
        legend_title="Legend",
        autosize=False,
        width=1000,
        height=600,
        margin={"l": 50, "r": 50, "b": 100, "t": 100, "pad": 4},
        plot_bgcolor="rgba(0,0,0,0)",
        paper_bgcolor="rgba(0,0,0,0)",
    )

    # Show the figure
    if show:
        fig.show()

    return fig


def plot_plot_ly(data: Reader, figure: NamedFigure, show: bool = False) -> plotlyFigure:

    fig = make_subplots(specs=[[{"secondary_y": True}]])

    for axis, figure_axis in enumerate([figure.left_axis, figure.right_axis]):
        for line in figure_axis.lines:
            try:
                variable = data.varNames(f"{line.key}$")[0]
                line_data = pd.DataFrame(data.values(variable))
            except (KeyError, IndexError):
                logger.warning(f"Key {line.key} not found in data")
                continue

            if line_data.empty:
                continue

            fig.add_trace(
                go.Scatter(
                    x=line_data.loc[0],
                    y=line_data.loc[1],
                    mode="lines",
                    name=line.label,
                ),
                secondary_y=bool(axis),
            )
    if not fig.data:
        return None
    fig.update_layout(
        xaxis_title="Simulation time [-]",
        yaxis_title=figure.left_axis.label,
        yaxis2_title=figure.right_axis.label,
        legend_title="Legend",
        autosize=False,
        width=1000,
        height=600,
        margin={"l": 50, "r": 50, "b": 100, "t": 100, "pad": 4},
    )

    # Show the figure
    if show:
        fig.show()

    return fig


def plot_element(
    data: Reader,
    element: BaseElement,
    plot_function: Callable[[Reader, NamedFigure], plotlyFigure | pyFigure] = plot,
) -> List[pyFigure]:
    figures = []
    subsystems = ["control", "emissions", "ventilation_inlets", "ventilation_outlets"]
    for figure in element.figures:
        fig = plot_function(data, figure)
        if fig:
            figures.append(fig)
        for subsystem in subsystems:
            if hasattr(element, subsystem) and getattr(element, subsystem) is not None:
                sub_element = getattr(element, subsystem)
                if isinstance(sub_element, list):
                    for sub in sub_element:
                        for figure_ in sub.figures:
                            fig = plot_function(data, figure_)
                            figures.append(fig)
                else:
                    for control_figure in sub_element.figures:
                        fig = plot_function(data, control_figure)
                        figures.append(fig)
    return figures


def add_element_figures(document: Document, data: Reader, element: BaseElement) -> None:
    # TODO: duplicate with the one bellow  !!! to be checked!!!
    figures = plot_element(data, element)
    for figure in figures:
        add_figure(document, figure, size=6)


def add_figure(document: Document, fig: pyFigure, size: int = 6) -> None:
    global FIGURE_COUNT  # noqa: PLW0603
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break()
    with tempfile.NamedTemporaryFile(suffix=".jpg") as f:
        fig.savefig(f.name, dpi=1200)
        # Add a figure
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(f.name, width=Inches(size))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the image
        caption = document.add_paragraph(
            f"Figure {FIGURE_COUNT}: test", style="Caption"
        )
        FIGURE_COUNT += 1
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break()
