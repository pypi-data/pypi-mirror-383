from pathlib import Path
from typing import Any, Dict, List, Union

import docx
from buildingspy.io.outputfile import Reader  # type: ignore
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from docx.table import _Cell

from trano.plot.plot import add_element_figures
from trano.reporting.reporting import ModelDocumentation
from trano.reporting.types import BaseNestedTable, Topic
from trano.reporting.utils import get_description

COLUMN_SIZE_WITH_DESCRIPTION = 3

TABLE_COUNT = 1


def _round(value: float | str | None) -> str:
    try:
        return str(round(value, 4))  # type: ignore
    except TypeError:
        return str(value)


def _set_color(cell: _Cell, key: str) -> None:
    if key == "type":
        cell._tc.get_or_add_tcPr().append(
            parse_xml(rf'<w:shd {nsdecls("w")} w:fill="#D3D3D3"/>')
        )
    if key == "name":
        cell._tc.get_or_add_tcPr().append(
            parse_xml(rf'<w:shd {nsdecls("w")} w:fill="#FF7F7F"/>')
        )
    if key == "thickness":
        cell._tc.get_or_add_tcPr().append(
            parse_xml(rf'<w:shd {nsdecls("w")} w:fill="#FFFFE0"/>')
        )


def create_table(doc: Document, data: Dict[str, Any]) -> None:  # type: ignore
    size = 2
    description = get_description()
    if any(key in description for key in data):
        size = COLUMN_SIZE_WITH_DESCRIPTION
    table = doc.add_table(cols=size, rows=len(data))  # type: ignore
    for i, (key, value) in enumerate(data.items()):
        table.rows[i].cells[0].text = key
        _set_color(table.rows[i].cells[0], key)
        if isinstance(value, dict):
            create_table(table.rows[i].cells[1], value)
        elif isinstance(value, list) and all(isinstance(v, dict) for v in value):
            for v in value:
                create_table(table.rows[i].cells[1], v)
        else:
            _set_color(table.rows[i].cells[1], key)
            table.rows[i].cells[1].text = _round(value)
            if size == COLUMN_SIZE_WITH_DESCRIPTION:
                table.rows[i].cells[2].text = str(description.get(key, "N/A"))


def add_table_caption(doc: Document, caption_text: str) -> None:  # type: ignore
    global TABLE_COUNT  # noqa: PLW0603
    caption = doc.add_paragraph(f"Table {TABLE_COUNT}: {caption_text}", style="Caption")  # type: ignore
    TABLE_COUNT += 1
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.space_after = Pt(0)


def create_tables_and_figures(
    doc: Document,  # type: ignore
    data: Union[List[BaseNestedTable], BaseNestedTable],
    topic: Topic,
    documentation: ModelDocumentation,
) -> None:
    if isinstance(data, list):
        for d in data:
            if d.name:
                add_table_caption(doc, f"Characteristics of {topic} {d.name}.")
                create_table(doc, d.model_dump())
                doc.add_paragraph()  # type: ignore
                insert_figure(doc, d.name, documentation)
                doc.add_paragraph()  # type: ignore

    elif data.name:
        add_table_caption(doc, f"Characteristics of {topic} {data.name}.")
        create_table(doc, data.model_dump())
        doc.add_paragraph()  # type: ignore
        insert_figure(doc, data.name, documentation)
        doc.add_paragraph()  # type: ignore


def insert_figure(doc: Document, element_name: str, documentation: ModelDocumentation) -> None:  # type: ignore
    if documentation.result is None:
        return
    mat = Reader(documentation.result.path, documentation.result.type)
    elements = [node for node in documentation.elements if node.name == element_name]
    if elements:
        element = elements[0]
        add_element_figures(doc, mat, element)


def to_docx(documentation: ModelDocumentation, path: Path) -> docx.document.Document:
    document = Document()
    document.add_page_break()  # type: ignore
    for doc in [
        documentation.spaces,
        documentation.systems,
        documentation.constructions,
    ]:
        document.add_heading(doc.topic, level=2)  # type: ignore
        if doc.introduction:
            document.add_paragraph(doc.introduction)
        create_tables_and_figures(document, doc.table, doc.topic, documentation)  # type: ignore
        if doc.conclusions:
            document.add_paragraph(doc.conclusions)
        document.add_page_break()  # type: ignore
    document.save(str(path))
    return document
