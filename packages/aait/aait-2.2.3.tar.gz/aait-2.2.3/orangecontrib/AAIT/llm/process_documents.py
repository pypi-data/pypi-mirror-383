import os
import json
import ntpath
from pathlib import Path

from Orange.data import Table, Domain, StringVariable

import fitz
import docx


def process_documents(dirpath):
    if dirpath is None or not Path(dirpath).exists():
        return None, None

    # Normalize dirpath
    dirpath = Path(dirpath).resolve()

    # get path from user selection
    embeddings = check_for_embeddings(dirpath)

    # Set selected path in the saved embeddings
    if embeddings is not None:
        common_path = Path(find_common_root(embeddings)).resolve()
        for row in embeddings:
            row_path = Path(str(row["path"].value)).resolve()
            # Replace common root with the current dirpath
            try:
                row["path"] = str(dirpath / row_path.relative_to(common_path))
            except ValueError:
                # If relative_to fails (paths not matching), just normalize
                row["path"] = str(row_path)

    # Verify which files are already processed
    files_to_process = get_files_to_process(dirpath, embeddings)

    rows = []
    for file in files_to_process:
        file = Path(file).resolve()
        content = extract_text(str(file))   # extractor may expect string
        filename = file.name
        row = [str(file), filename, content]  # store strings in Orange table
        rows.append(row)

    # Build a table with the constructed rows
    path_var = StringVariable("path")
    name_var = StringVariable("name")
    content_var = StringVariable("content")
    domain = Domain(attributes=[], metas=[path_var, name_var, content_var])
    out_data = Table.from_list(domain=domain, rows=rows)
    return out_data, embeddings


def check_for_embeddings(folder_path):
    """
    Check for an embeddings.pkl file in a given folder. Return its content if it exists.

    Parameters:
        folder_path (str | Path): The path to the folder where embeddings.pkl may exist.

    Returns:
        Table or None: The content of embeddings.pkl, or None if not found.
    """
    folder_path = Path(folder_path).resolve()

    filepaths = [
        folder_path / "embeddings_question.pkl",
        folder_path / "embeddings.pkl"
    ]

    for filepath in filepaths:
        if filepath.exists():
            return Table.from_file(str(filepath))  # Table.from_file expects a str
    return None


def find_common_root(data_table, column_name="path"):
    """Finds the common root path from a column of file paths in an Orange Data Table."""
    paths = [Path(str(row[column_name].value)).resolve()
             for row in data_table if row[column_name] is not None]
    if not paths:
        return ""
    return str(Path(os.path.commonpath(paths)).resolve())


def get_files_to_process(folder_path, table=None):
    """
    Finds all PDF/DOCX files in a folder (including subfolders) that are not already in the table
    or that have changed since last check (based on file size).

    :param folder_path: Path to the folder to scan for documents.
    :param table: Orange Data Table with column "path".
    :return: List of new/changed file paths.
    """
    supported_extensions = {".pdf", ".docx"}

    folder_path = Path(folder_path).resolve()
    filepath_sizes = folder_path / "sizes.json"

    print(filepath_sizes)
    print(filepath_sizes.exists())

    # Load previous file sizes
    if filepath_sizes.exists():
        with open(filepath_sizes, "r") as json_file:
            sizes = {Path(k): v for k, v in json.load(json_file).items()}
    else:
        sizes = {}

    print(sizes)

    # Extract the existing paths from the Orange Data Table
    if table:
        # Orange stores metas as strings → turn them into Paths
        existing_paths = {Path(str(p)).resolve() for p in table[:, "path"].metas.flatten()}
    else:
        existing_paths = set()

    new_files = []

    # Walk through the folder and subfolders
    for file in folder_path.rglob("*"):
        if file.suffix.lower() in supported_extensions:
            file = file.resolve()
            size = file.stat().st_size

            if file not in existing_paths:
                # New file
                new_files.append(str(file))
                sizes[Path(ntpath.basename(file))] = size
            else:
                # File already in table: check if size changed
                print("File: ", Path(ntpath.basename(file)))
                old_size = sizes.get(Path(ntpath.basename(file)))
                print(old_size)
                if old_size is None or old_size != size:
                    new_files.append(str(file))
                    table = remove_from_table(file, table)
                    sizes[Path(ntpath.basename(file))] = size


    # Save updated sizes.json (keys must be strings for JSON)
    with open(filepath_sizes, "w") as json_file:
        json.dump({str(k): v for k, v in sizes.items()}, json_file, indent=4)

    return new_files


def remove_from_table(filepath, table):
    """
    Remove rows from the Orange table where 'path' matches the given filepath.
    """
    filepath = Path(filepath).resolve()

    filtered_table = Table.from_list(
        domain=table.domain,
        rows=[
            row for row in table
            if Path(str(row["path"].value)).resolve() != filepath
        ]
    )
    return filtered_table


def extract_text(filepath):
    """
    Extrait le texte d'un fichier en fonction de son type (PDF ou DOCX).

    :param filepath: Chemin vers le fichier.
    :return: Texte extrait du fichier sous forme de chaîne.
    """
    try:
        # Vérifie l'extension du fichier
        file_extension = os.path.splitext(filepath)[1].lower()

        if file_extension == ".pdf":
            return extract_text_from_pdf(filepath)
        elif file_extension == ".docx":
            return extract_text_from_docx(filepath)
        else:
            raise ValueError("Format de fichier non supporté. Utilisez un fichier PDF ou DOCX.")
    except Exception as e:
        print(f"Erreur lors de l'extraction de texte depuis {filepath}: {e}")
        return "Extraction Error"


def extract_text_from_pdf(pdf_path):
    """
    Extrait le texte d'un fichier PDF.

    :param pdf_path: Chemin vers le fichier PDF.
    :return: Texte extrait du PDF sous forme de chaîne.
    """
    try:
        # Ouvre le fichier PDF
        pdf_document = fitz.open(pdf_path)
        extracted_text = ""

        # Parcourt toutes les pages et extrait le texte
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            extracted_text += page.get_text()

        pdf_document.close()
        return extracted_text
    except Exception as e:
        print(f"Erreur lors de l'extraction de texte depuis {pdf_path}: {e}")
        return "Extraction Error"


def extract_text_from_docx(docx_path):
    """
    Extrait le texte d'un fichier DOCX en conservant l'ordre des éléments (paragraphes, tableaux et titres).

    :param docx_path: Chemin vers le fichier DOCX.
    :return: Texte extrait du document sous forme de chaîne.
    """
    try:
        doc = docx.Document(docx_path)
        extracted_text = []
        title_numbers = {}  # Dictionary to track numbering per heading level

        for para in doc.paragraphs:
            # Vérifie si c'est un titre
            if para.style.name.startswith('Heading'):
                heading_level = int(para.style.name.split()[-1])  # Niveau du titre (1, 2, 3, etc.)
                heading_text = para.text.strip()

                # Met à jour la numérotation des titres
                if heading_level not in title_numbers:
                    title_numbers[heading_level] = 1  # Nouveau niveau
                else:
                    title_numbers[heading_level] += 1  # Incrémente niveau actuel

                # Réinitialise les niveaux inférieurs
                for level in list(title_numbers.keys()):
                    if level > heading_level:
                        del title_numbers[level]

                # Forme le numéro du titre (ex: "1", "1.1", "1.2.1")
                full_title = ".".join(str(title_numbers[i]) for i in sorted(title_numbers.keys()))
                extracted_text.append(f"\n{full_title} {heading_text}")  # Ajoute le titre formaté
            else:
                extracted_text.append(para.text.strip())  # Ajoute le paragraphe

        # Parcourt les tableaux du document
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]  # Extrait le texte de chaque cellule
                table_text.append("\t".join(row_text))  # Sépare les colonnes par des tabulations
            extracted_text.append("\n".join(table_text))  # Ajoute le tableau sous forme de texte

        return "\n".join(filter(None, extracted_text))  # Retourne le texte en filtrant les vides

    except Exception as e:
        print(f"Erreur lors de l'extraction de texte depuis {docx_path}: {e}")
        return "Extraction Error"



def get_pages_of_extract(pdf_path, extract):
    """
    Identify the pages that a given extract belongs to.

    :param pdf_path: The path of the pdf to search in.
    :param extract: The text snippet to locate.
    :return: A list of page numbers the extract spans.
    """
    full_text, page_mapping = load_pdf_with_sparse_mapping(pdf_path)
    # Find the start index of the extract in the full text
    start_index = full_text.find(extract)
    if start_index == -1:
        return []  # Extract not found

    # Determine the end index of the extract
    end_index = start_index + len(extract) - 1

    # Find all pages the extract spans
    pages = []
    for page, (start, end) in page_mapping.items():
        if start <= end_index and end >= start_index:
            pages.append(page)

    if pages == []:
        return [1]
    return pages


def load_pdf_with_sparse_mapping(pdf_path):
    doc = fitz.open(pdf_path)
    full_text = ""
    page_mapping = {}  # Sparse mapping: {page_num: (start_index, end_index)}

    for page_num in range(len(doc)):
        page_text = doc[page_num].get_text()
        start_index = len(full_text)
        full_text += page_text
        end_index = len(full_text) - 1
        page_mapping[page_num + 1] = (start_index, end_index)

    doc.close()
    return full_text, page_mapping
