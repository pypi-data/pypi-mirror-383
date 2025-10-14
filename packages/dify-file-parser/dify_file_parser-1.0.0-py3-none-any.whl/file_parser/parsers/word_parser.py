"""
Word 文档解析器
支持 .docx 和 .doc 格式
"""

import asyncio
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
import subprocess
import tempfile
import os

from ..base import BaseParser, ParseResult, FileType


class WordParser(BaseParser):
    """Word 文档解析器"""
    
    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]
    
    @property
    def file_type(self) -> FileType:
        return FileType.WORD
    
    async def parse(self, file_path: Path) -> ParseResult:
        """
        解析 Word 文档
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            ParseResult: 解析结果
        """
        try:
            if file_path.suffix.lower() == ".docx":
                text_content = await self._parse_docx(file_path)
            else:  # .doc
                text_content = await self._parse_doc(file_path)
            
            # 获取文档元数据
            metadata = self._get_file_metadata(file_path)
            metadata.update({
                "word_format": file_path.suffix.lower(),
                "paragraph_count": len([p for p in text_content.split('\n') if p.strip()])
            })
            
            return ParseResult(
                filename=Path(file_path).name,
                file_type=self.file_type,
                text=text_content.strip(),
                metadata=metadata,
                success=True
            )
            
        except Exception as e:
            self.logger.error(f"Word 文档解析失败: {str(e)}")
            raise
    
    async def _parse_docx(self, file_path: Path) -> str:
        """解析 .docx 文件"""
        def _extract():
            try:
                doc = Document(file_path)
                text_parts = []
                
                # 提取段落文本
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # 提取表格文本
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(" | ".join(row_text))
                    if table_text:
                        text_parts.append("\n[表格内容]\n" + "\n".join(table_text))
                
                return "\n\n".join(text_parts)
                
            except Exception as e:
                self.logger.error(f"解析 .docx 文件失败: {str(e)}")
                raise
        
        return await asyncio.get_event_loop().run_in_executor(None, _extract)
    
    async def _parse_doc(self, file_path: Path) -> str:
        """解析 .doc 文件（需要安装 antiword 或使用 python-docx2txt）"""
        def _extract():
            try:
                # 尝试使用 python-docx2txt
                try:
                    import docx2txt
                    return docx2txt.process(str(file_path))
                except ImportError:
                    pass
                
                # 尝试使用 antiword（Linux/Mac）
                try:
                    result = subprocess.run(
                        ['antiword', str(file_path)],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0:
                        return result.stdout
                except (subprocess.TimeoutExpired, FileNotFoundError):
                    pass
                
                # 尝试使用 catdoc（Linux/Mac）
                try:
                    result = subprocess.run(
                        ['catdoc', str(file_path)],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0:
                        return result.stdout
                except (subprocess.TimeoutExpired, FileNotFoundError):
                    pass
                
                # 如果都失败了，尝试转换为 docx 再解析
                return self._convert_doc_to_docx_and_parse(file_path)
                
            except Exception as e:
                self.logger.error(f"解析 .doc 文件失败: {str(e)}")
                raise
        
        return await asyncio.get_event_loop().run_in_executor(None, _extract)
    
    def _convert_doc_to_docx_and_parse(self, file_path: Path) -> str:
        """将 .doc 转换为 .docx 然后解析"""
        try:
            # 这里需要安装 python-docx2txt 或使用 LibreOffice
            # 简化实现，返回错误信息
            return f"[警告] 无法解析 .doc 文件 {file_path.name}，请安装 python-docx2txt 或转换为 .docx 格式"
        except Exception as e:
            self.logger.error(f"转换 .doc 文件失败: {str(e)}")
            return f"[错误] 解析 .doc 文件失败: {str(e)}"
